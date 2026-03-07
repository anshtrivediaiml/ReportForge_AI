"""
Builder Agent - Assembles final DOCX document
"""
import json
import base64
import zlib
import re
from pathlib import Path
from typing import Dict, Any
from loguru import logger
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.docx_generator import DOCXGenerator
import requests


class BuilderAgent:
    """Builds the final DOCX document from content and guidelines"""
    
    def __init__(self, output_dir: str = "outputs/final", job_id: str = None):
        # Create job-specific output directory to prevent content mixing
        if job_id:
            self.output_dir = Path(output_dir) / f"job_{job_id}"
        else:
            self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.job_id = job_id
        # Track generated diagrams to prevent duplicates (per job)
        self.generated_diagrams = {}  # {normalized_mermaid_code: section_num}
        # Track skipped/failed components for summary report
        self.skipped_components = {
            "diagrams": [],
            "tables": [],
            "subsections": []
        }
    
    def build_document(
        self,
        content: Dict[str, Any],
        guidelines: Dict[str, Any],
        outline: Dict[str, Any],
        output_filename: str = "Technical_Report.docx"
    ) -> str:
        """
        Build complete DOCX document
        
        Args:
            content: Generated content from WriterAgent
            guidelines: Formatting rules from ParserAgent
            outline: Report structure from PlannerAgent
            output_filename: Output file name
            
        Returns:
            Path to generated document
        """
        logger.info("=== BUILDER AGENT STARTED ===")
        logger.info("Building DOCX document...")
        
        # Create document
        doc = DOCXGenerator(guidelines)
        
        # Initialize tracking for figures and tables
        doc.figures_list = []
        doc.tables_list = []
        
        # Add title page
        self._add_title_page(doc, content)
        doc.add_page_break()
        
        # Add table of contents (before chapters)
        doc.add_table_of_contents(outline)
        doc.add_page_break()  # Page break after TOC
        
        # First pass: Collect all figures and tables by processing chapters
        temp_doc = DOCXGenerator(guidelines)
        temp_doc.figures_list = []
        temp_doc.tables_list = []
        
        chapters = content.get("chapters", [])
        for i, chapter in enumerate(chapters):
            self._add_chapter(temp_doc, chapter, collect_only=True)
        
        # Now add List of Figures and List of Tables before Chapter 1
        if temp_doc.figures_list:
            doc.figures_list = temp_doc.figures_list
            self._add_list_of_figures(doc)
            doc.add_page_break()
        
        if temp_doc.tables_list:
            doc.tables_list = temp_doc.tables_list
            self._add_list_of_tables(doc)
            doc.add_page_break()
        
        # Second pass: Add all chapters normally (they'll populate lists again, but that's okay)
        for i, chapter in enumerate(chapters):
            self._add_chapter(doc, chapter)
            # Add page break after each chapter (including the last one, before references)
            doc.add_page_break()
        
        # Add references section (AI-generated)
        # Extract codebase structure from content if available
        codebase_info = {
            "name": content.get("report_title", "Technical Report"),
            "technologies": []
        }
        # Try to extract from chapters
        for chapter in content.get("chapters", []):
            for section in chapter.get("sections", []):
                content_text = section.get("content", "").lower()
                if "python" in content_text:
                    codebase_info["technologies"].append("Python")
                if "javascript" in content_text or "node" in content_text:
                    codebase_info["technologies"].append("JavaScript")
                if "api" in content_text or "rest" in content_text:
                    codebase_info["technologies"].append("API")
                if "database" in content_text or "sql" in content_text:
                    codebase_info["technologies"].append("Database")
        
        self._add_references_section(doc, content, outline, codebase_info)
        # Don't add page break after references (last section)
        
        # Add headers/footers
        doc.add_header_footer()
        
        # Save document
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        # Log summary of skipped components
        self._log_skipped_components_summary()
        
        logger.success("=== BUILDER AGENT COMPLETED ===")
        logger.info(f"Document saved: {output_path}")
        
        return str(output_path)
    
    def _add_title_page(self, doc: DOCXGenerator, content: Dict[str, Any]):
        """Add title page"""
        logger.info("Adding title page...")
        
        # Title
        title = content.get("report_title", "Technical Report")
        # HARDENING: Validate title
        title = self._validate_heading(title)
        if not title:
            title = "Technical Report"  # Fallback
        doc.add_chapter_heading(0, title)
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Subtitle
        doc.add_paragraph("Technical Documentation")
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Author info
        doc.add_paragraph("Prepared by: Student Name")
        doc.add_paragraph("Student ID: 22AIML056")
        doc.add_paragraph("")
        
        # Institution
        doc.add_paragraph("CSPIT")
        doc.add_paragraph("Department of Artificial Intelligence & Machine Learning")
        
        logger.success("Title page added")
    
    def _add_chapter(self, doc: DOCXGenerator, chapter: Dict[str, Any], collect_only: bool = False):
        """Add a complete chapter with sections"""
        # Determine if this is a conclusion chapter (needed for both collect_only and normal modes)
        chapter_title = chapter.get("chapter_title", "Chapter")
        # Fix: Handle escaped JSON strings for chapter titles
        if isinstance(chapter_title, str):
            chapter_title = chapter_title.strip('"\'')
        # HARDENING: Validate chapter title
        chapter_title = self._validate_heading(chapter_title)
        if not chapter_title:
            chapter_title = "Chapter"  # Fallback
        is_conclusion = "conclusion" in chapter_title.strip().upper()
        
        if not collect_only:
            chapter_num = chapter.get("chapter_number", 1)

            if is_conclusion:
                logger.info("Adding Conclusion")
                # Conclusion should not be labelled with a chapter number in the heading
                doc.add_conclusion_heading(chapter_title)
            else:
                logger.info(f"Adding Chapter {chapter_num}: {chapter_title}")
                doc.add_chapter_heading(chapter_num, chapter_title)
        
        # Add sections
        sections = chapter.get("sections", [])
        if not isinstance(sections, list):
            logger.warning(f"Sections is not a list (type: {type(sections)}), skipping")
            sections = []
        
        for section in sections:
            # Ensure section is a dict
            if not isinstance(section, dict):
                logger.warning(f"Section is not a dict (type: {type(section)}), skipping")
                continue
            self._add_section(doc, section, is_conclusion=is_conclusion, collect_only=collect_only)
        
        if not collect_only:
            logger.success(f"Chapter {chapter.get('chapter_number', 1)} added with {len(sections)} sections")
    
    def _add_section(self, doc: DOCXGenerator, section: Dict[str, Any], is_conclusion: bool = False, collect_only: bool = False):
        """Add a section with content, tables, and diagrams"""
        # Validate section is a dict
        if not isinstance(section, dict):
            logger.error(f"Section must be a dict, got {type(section)}, skipping")
            return
        
        # Support both "number" and "section_number" keys
        section_num = section.get("number") or section.get("section_number") or ""
        # Handle escaped JSON strings (double-quoted strings)
        if isinstance(section_num, str):
            section_num = section_num.strip('"')
        section_title = section.get("title", "Section")
        # Handle escaped JSON strings
        if isinstance(section_title, str):
            section_title = section_title.strip('"')
        content = section.get("content", "")
        # Handle escaped JSON strings
        if isinstance(content, str):
            # Remove outer quotes if present
            content = content.strip('"')
            # Unescape newlines and other escape sequences
            content = content.replace('\\n', '\n').replace('\\"', '"')

        # Remove any embedded Mermaid code blocks from the visible text
        content = re.sub(r"```mermaid[\s\S]*?```", "", content, flags=re.MULTILINE)
        
        # HARDENING: Validate and sanitize section title
        section_title = self._validate_heading(section_title)
        if not section_title:
            section_title = "Section"  # Fallback
        
        logger.debug(f"Adding section {section_num}: {section_title}")
        
        if not collect_only:
            # Add section heading; for conclusion sections, always use empty number
            if is_conclusion:
                section_num = ""  # Force empty for conclusion sections
            doc.add_section_heading(section_num, section_title, is_conclusion=is_conclusion)
            
            # Add content paragraphs (handle bullet points and lists, split long paragraphs)
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if para:
                    # Check if paragraph contains bullet points or numbered lists
                    lines = para.split("\n")
                    if len(lines) > 1 and any(line.strip().startswith(("•", "-", "*")) or re.match(r'^\d+\.', line.strip()) for line in lines):
                        # Handle as list - add each line as a list item
                        for line in lines:
                            line = line.strip()
                            if line:
                                # Remove bullet markers and add as paragraph with indent
                                clean_line = re.sub(r'^[•\-\*\d+\.]\s*', '', line)
                                if clean_line:
                                    # HARDENING: Sanitize text coherence and enforce epistemic boundary
                                    clean_line = self._sanitize_text_coherence(clean_line)
                                    clean_line = self._enforce_epistemic_boundary(clean_line)
                                    if clean_line.strip():
                                        p = doc.document.add_paragraph(clean_line, style='List Bullet')
                    else:
                        # Regular paragraph - split if too long (more than 150 words or 800 characters)
                        # HARDENING: Sanitize text coherence and enforce epistemic boundary
                        para = self._sanitize_text_coherence(para)
                        para = self._enforce_epistemic_boundary(para)
                        
                        word_count = len(para.split())
                        if word_count > 150 or len(para) > 800:
                            # Split long paragraph at sentence boundaries
                            split_paras = self._split_long_paragraph(para)
                            for split_para in split_paras:
                                if split_para.strip():
                                    doc.add_paragraph(split_para.strip())
                        else:
                            doc.add_paragraph(para)
        
        # Add table if provided (track for list even in collect_only mode)
        table_data = section.get("table_data")
        # FIX: Handle escaped JSON strings (table_data might be a JSON string)
        if isinstance(table_data, str):
            try:
                table_data = json.loads(table_data)
                logger.debug(f"Parsed table_data JSON string for section {section_num}")
            except (json.JSONDecodeError, ValueError) as e:
                reason = f"Failed to parse JSON: {str(e)[:50]}"
                logger.warning(f"Failed to parse table_data JSON for section {section_num}, skipping table")
                if not collect_only:  # Only track in normal pass to avoid duplicates
                    self.skipped_components["tables"].append((section_num, reason))
                table_data = None
        if table_data and isinstance(table_data, list):
            # HARDENING: Validate and sanitize table data
            validated_table = self._validate_table_data(table_data)
            if validated_table:
                table_data = validated_table
            else:
                # Table validation emptied the table
                reason = "Table validation removed all invalid cells"
                logger.warning(f"Table validation failed for section {section_num}, skipping")
                if not collect_only:  # Only track in normal pass
                    self.skipped_components["tables"].append((section_num, reason))
                table_data = None
            
            if table_data:  # Only proceed if validation didn't empty the table
                # Generate table label (e.g., "Table 2.1")
                table_label = section.get("table_label") or f"Table {section_num}" if section_num else "Table"
                table_title = section.get("table_title") or f"Summary for {section_title}"
                # HARDENING: Validate table title
                table_title = self._validate_heading(table_title)
                if not table_title:
                    table_title = "Summary"
                
                if collect_only:
                    # Just track, don't add to document
                    if not hasattr(doc, 'tables_list'):
                        doc.tables_list = []
                    doc.tables_list.append({
                        "label": table_label,
                        "title": table_title
                    })
                else:
                    doc.add_table(table_data, title=table_title, table_label=table_label)
        
        # Add diagram only if Mermaid code provided (generate via multiple services)
        mermaid_code = section.get("mermaid_code")
        if mermaid_code:
            # Fix: Handle quoted/escaped mermaid code
            if isinstance(mermaid_code, str):
                mermaid_code = mermaid_code.strip('"\'')
                # Fix escaped newlines
                mermaid_code = mermaid_code.replace('\\n', '\n')
            
            # Clean mermaid code - remove markdown code blocks if present
            mermaid_code = mermaid_code.strip()
            if mermaid_code.startswith("```"):
                # Extract code from markdown block
                lines = mermaid_code.split("\n")
                mermaid_code = "\n".join([l for l in lines if not l.strip().startswith("```")])
            
            # Fix: Validate mermaid code completeness before processing
            if mermaid_code:
                # Check for balanced brackets (basic completeness check)
                open_brackets = mermaid_code.count('[')
                close_brackets = mermaid_code.count(']')
                open_parens = mermaid_code.count('(')
                close_parens = mermaid_code.count(')')
                
                # Check if code appears incomplete (truncated)
                if open_brackets != close_brackets or open_parens != close_parens:
                    reason = "Incomplete code (unbalanced brackets)"
                    logger.warning(f"Incomplete mermaid code for section {section_num} (unbalanced brackets), skipping")
                    if not collect_only:  # Only track in normal pass to avoid duplicates
                        self.skipped_components["diagrams"].append((section_num, reason))
                    mermaid_code = None
                # Check minimum length and basic structure
                elif len(mermaid_code.strip()) < 10:
                    logger.debug(f"Skipping empty/invalid mermaid code for section {section_num}")
                    mermaid_code = None
                # Check if code has valid diagram type declaration
                elif not any(mermaid_code.strip().lower().startswith(diag_type) for diag_type in 
                           ['graph', 'flowchart', 'sequencediagram', 'classdiagram', 'erdiagram']):
                    logger.warning(f"Mermaid code missing valid diagram type for section {section_num}, skipping")
                    mermaid_code = None
            
            if mermaid_code:
                # HARDENING: Validate diagram semantics before generating
                validated_code, is_valid = self._validate_diagram_semantics(mermaid_code)
                if not is_valid or not validated_code:
                    reason = "Invalid semantics (outcome/error-based diagram)"
                    logger.warning(f"Diagram semantics invalid for section {section_num}, skipping")
                    if not collect_only:  # Only track in normal pass to avoid duplicates
                        self.skipped_components["diagrams"].append((section_num, reason))
                else:
                    fig_label = section.get("figure_label") or f"Fig {section_num}"
                    fig_desc = section.get("figure_desc") or f"Diagram for {section_title}"
                    # Fix: Handle escaped JSON strings for figure labels/descriptions
                    if isinstance(fig_label, str):
                        fig_label = fig_label.strip('"\'')
                    if isinstance(fig_desc, str):
                        fig_desc = fig_desc.strip('"\'')
                    # HARDENING: Validate figure description
                    fig_desc = self._sanitize_text_coherence(fig_desc)
                    fig_label = self._validate_heading(fig_label)
                    if not fig_label:
                        fig_label = f"Fig {section_num}"
                    
                    logger.info(f"Attempting to generate diagram for section {section_num}")
                    result = self._generate_diagram(validated_code, section_num)
                
                if result == "DUPLICATE":
                    # Duplicate diagram - track for summary
                    reason = "Duplicate diagram (similar structure already generated)"
                    logger.debug(f"Skipping duplicate diagram for section {section_num} (no placeholder shown)")
                    if not collect_only:  # Only track in normal pass
                        self.skipped_components["diagrams"].append((section_num, reason))
                elif result and isinstance(result, str):
                    # Successfully generated diagram
                    logger.success(f"Diagram generated successfully for section {section_num}")
                    if collect_only:
                        # Just track, don't add to document
                        if not hasattr(doc, 'figures_list'):
                            doc.figures_list = []
                        doc.figures_list.append({
                            "label": fig_label,
                            "description": fig_desc
                        })
                    else:
                        doc.add_figure(result, fig_label, fig_desc)
                else:
                    # Generation failed - only show placeholder if it was a genuine attempt (not duplicate)
                    # Check if it was marked as duplicate in the warning
                    logger.debug(f"Diagram generation failed for section {section_num}, skipping (no placeholder)")
                    # Don't show placeholder for failed generations - if LLM thought it needed a diagram but it failed,
                    # it's better to skip it than show a placeholder
        
        # Add subsections if present
        subsections = section.get("subsections", [])
        # FIX: Handle escaped JSON strings (subsections might be a JSON string)
        if isinstance(subsections, str):
            try:
                subsections = json.loads(subsections)
                logger.debug(f"Parsed subsections JSON string for section {section_num}")
            except (json.JSONDecodeError, ValueError) as e:
                reason = f"Failed to parse JSON: {str(e)[:50]}"
                logger.warning(f"Failed to parse subsections JSON for section {section_num}, skipping subsections")
                if not collect_only:  # Only track in normal pass to avoid duplicates
                    self.skipped_components["subsections"].append((section_num, reason))
                subsections = []
        if subsections:
            # Ensure subsections is a list
            if not isinstance(subsections, list):
                reason = f"Invalid data type: {type(subsections).__name__}"
                logger.warning(f"Subsections is not a list (type: {type(subsections)}), skipping")
                if not collect_only:  # Only track in normal pass to avoid duplicates
                    self.skipped_components["subsections"].append((section_num, reason))
                subsections = []
            
            # Generate proper subsection numbers based on section number
            for idx, subsection in enumerate(subsections, 1):
                # Ensure subsection is a dict, not a string
                if not isinstance(subsection, dict):
                    logger.warning(f"Subsection is not a dict (type: {type(subsection)}), converting to dict")
                    # Convert string to dict
                    if isinstance(subsection, str):
                        subsection = {"title": "Subsection", "content": subsection}
                    else:
                        logger.error(f"Cannot convert subsection type {type(subsection)} to dict, skipping")
                        continue
                
                # Generate proper subsection number (e.g., "2.1.1", "2.1.2", "3.2.1")
                if section_num:
                    subsection_num = f"{section_num}.{idx}"
                else:
                    subsection_num = f"{idx}.{idx}.{idx}"  # Fallback if no section number
                
                self._add_subsection(doc, subsection, subsection_num)
    
    def _add_subsection(self, doc: DOCXGenerator, subsection: Dict[str, Any], subsection_num: str):
        """Add a subsection with proper numbering"""
        # Validate subsection is a dict
        if not isinstance(subsection, dict):
            logger.error(f"Subsection must be a dict, got {type(subsection)}, skipping")
            return
        
        # Use provided subsection_num instead of defaulting to "1.1.1"
        subsection_title = subsection.get("title", "Subsection")
        content = subsection.get("content", "")
        
        # Fix: Handle escaped JSON strings
        if isinstance(subsection_title, str):
            subsection_title = subsection_title.strip('"\'')
        if isinstance(content, str):
            content = content.strip('"\'')
            content = content.replace('\\n', '\n').replace('\\"', '"')
        
        # HARDENING: Validate subsection title
        subsection_title = self._validate_heading(subsection_title)
        if not subsection_title:
            subsection_title = "Subsection"  # Fallback
        
        doc.add_subsection_heading(subsection_num, subsection_title)
        
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if para:
                # HARDENING: Sanitize text coherence and enforce epistemic boundary
                para = self._sanitize_text_coherence(para)
                para = self._enforce_epistemic_boundary(para)
                
                # Split long paragraphs for readability
                word_count = len(para.split())
                if word_count > 150 or len(para) > 800:
                    split_paras = self._split_long_paragraph(para)
                    for split_para in split_paras:
                        if split_para.strip():
                            doc.add_paragraph(split_para.strip())
                else:
                    doc.add_paragraph(para)
    
    def _split_long_paragraph(self, paragraph: str) -> list:
        """Split a long paragraph into 2-3 shorter paragraphs at natural break points"""
        # Target: 100-150 words per paragraph
        words = paragraph.split()
        if len(words) <= 150:
            return [paragraph]  # No need to split
        
        # Find sentence boundaries
        sentences = re.split(r'([.!?]\s+)', paragraph)
        # Recombine sentences with their punctuation
        combined_sentences = []
        for i in range(0, len(sentences) - 1, 2):
            if i + 1 < len(sentences):
                combined_sentences.append(sentences[i] + sentences[i + 1])
            else:
                combined_sentences.append(sentences[i])
        
        if len(combined_sentences) <= 1:
            # Can't split by sentences, try by commas or conjunctions
            return self._split_by_conjunctions(paragraph)
        
        # Group sentences into paragraphs of ~100-150 words
        result = []
        current_para = []
        current_word_count = 0
        
        for sentence in combined_sentences:
            sentence_words = len(sentence.split())
            if current_word_count + sentence_words > 150 and current_para:
                # Start new paragraph
                result.append(' '.join(current_para))
                current_para = [sentence]
                current_word_count = sentence_words
            else:
                current_para.append(sentence)
                current_word_count += sentence_words
        
        if current_para:
            result.append(' '.join(current_para))
        
        return result if len(result) > 1 else [paragraph]
    
    def _split_by_conjunctions(self, paragraph: str) -> list:
        """Split paragraph at conjunctions if sentence splitting didn't work"""
        # Split at common conjunctions that indicate natural breaks
        conjunctions = [r'\s+However,\s+', r'\s+Moreover,\s+', r'\s+Furthermore,\s+', 
                       r'\s+Additionally,\s+', r'\s+Consequently,\s+', r'\s+Therefore,\s+',
                       r'\s+In addition,\s+', r'\s+On the other hand,\s+', r'\s+That said,\s+']
        
        parts = [paragraph]
        for conj_pattern in conjunctions:
            new_parts = []
            for part in parts:
                splits = re.split(conj_pattern, part, maxsplit=1)
                if len(splits) > 1:
                    new_parts.extend(splits)
                else:
                    new_parts.append(part)
            parts = new_parts
            if len(parts) >= 2:
                break
        
        # If we have 2-3 parts and they're reasonable length, return them
        if 2 <= len(parts) <= 3:
            return [p.strip() for p in parts if p.strip()]
        
        # Otherwise, just split in the middle
        mid_point = len(paragraph) // 2
        # Find a good break point (space or punctuation near the middle)
        for i in range(mid_point - 50, mid_point + 50):
            if i < len(paragraph) and paragraph[i] in '.!?':
                return [paragraph[:i+1].strip(), paragraph[i+1:].strip()]
        
        # Last resort: split at middle space
        words = paragraph.split()
        mid_word = len(words) // 2
        return [' '.join(words[:mid_word]), ' '.join(words[mid_word:])]

    def _shorten_label(self, label: str, max_length: int = 15) -> str:
        """Create a readable, shortened label from a long name"""
        if len(label) <= max_length:
            return label
        
        # If it's camelCase or PascalCase, try to break it up
        import re
        # Split camelCase/PascalCase
        words = re.findall(r'[A-Z]?[a-z]+|[A-Z]+(?=[A-Z]|$)', label)
        if words and len(words) > 1:
            # Take first few words that fit
            result = []
            for word in words:
                if len(' '.join(result + [word])) <= max_length:
                    result.append(word)
                else:
                    break
            if result:
                return ' '.join(result)
        
        # If still too long, truncate intelligently
        if len(label) > max_length:
            # Try to truncate at word boundaries
            truncated = label[:max_length-3]
            last_space = truncated.rfind(' ')
            if last_space > max_length * 0.6:  # If we can keep at least 60% of the text
                return truncated[:last_space] + "..."
            return truncated + "..."
        
        return label[:max_length]
    
    def _clean_mermaid_code(self, mermaid_code: str) -> str:
        """Clean and validate Mermaid code - make it syntactically correct and readable"""
        if not mermaid_code:
            return ""
        
        import re
        
        # Remove markdown code blocks
        mermaid_code = mermaid_code.strip()
        if mermaid_code.startswith("```mermaid"):
            lines = mermaid_code.split("\n")
            mermaid_code = "\n".join([l for l in lines if not l.strip().startswith("```")])
        elif mermaid_code.startswith("```"):
            lines = mermaid_code.split("\n")
            mermaid_code = "\n".join([l for l in lines if not l.strip().startswith("```")])
        
        mermaid_code = mermaid_code.strip()
        
        # Remove any leading/trailing quotes or JSON escape characters
        mermaid_code = mermaid_code.strip('"\'')
        mermaid_code = mermaid_code.replace('\\n', '\n')
        mermaid_code = mermaid_code.replace('\\"', '"')
        
        # Remove empty lines
        lines = [l.strip() for l in mermaid_code.split("\n") if l.strip()]
        if not lines:
            return ""
        
        mermaid_code = "\n".join(lines)
        
        # Check if it starts with a valid diagram type - if so, preserve it mostly as-is
        valid_starts = ["graph", "flowchart", "sequenceDiagram", "classDiagram", "erDiagram", 
                       "gantt", "pie", "stateDiagram", "journey"]
        
        first_line_lower = lines[0].lower().strip() if lines else ""
        has_valid_start = any(first_line_lower.startswith(start) for start in valid_starts)
        
        if not has_valid_start:
            # Try to find valid start in lines
            for i, line in enumerate(lines):
                line_lower = line.lower().strip()
                if any(line_lower.startswith(start) for start in valid_starts):
                    mermaid_code = "\n".join(lines[i:])
                    lines = mermaid_code.split("\n")
                    break
                elif "graph" in line_lower or "flowchart" in line_lower:
                    # Add graph declaration if missing
                    mermaid_code = "graph TD\n" + "\n".join(lines[i+1:] if i+1 < len(lines) else [])
                    lines = mermaid_code.split("\n")
                    break
        
        # For sequence diagrams and class diagrams, preserve more of the original structure
        is_sequence = any(l.strip().lower().startswith("sequencediagram") for l in lines)
        is_class = any(l.strip().lower().startswith("classdiagram") for l in lines)
        
        if is_sequence or is_class:
            # For complex diagrams, do minimal cleaning - just fix obvious issues
            cleaned_lines = []
            for line in lines:
                line = line.rstrip()  # Preserve indentation
                if not line.strip():
                    continue
                # Only fix obvious syntax errors
                line = re.sub(r'->\s*->', '-->', line)
                line = re.sub(r'--\s*>', '-->', line)
                cleaned_lines.append(line)
            return "\n".join(cleaned_lines)
        
        # For simple graphs/flowcharts, do more aggressive cleaning
        cleaned_lines = []
        for line in lines:
            original_line = line
            line = line.strip()
            if not line:
                continue
            
            # Fix arrow syntax
            line = re.sub(r'->\s*->', '-->', line)
            line = re.sub(r'--\s*>', '-->', line)
            
            # Preserve valid lines
            if line.startswith(('graph', 'flowchart', 'sequenceDiagram', 'classDiagram', 'erDiagram', 'gantt', 'pie', 'stateDiagram', 'journey')):
                cleaned_lines.append(line)
            elif '-->' in line:
                # Connection line - clean node IDs but preserve structure
                parts = line.split('-->')
                if len(parts) == 2:
                    left = parts[0].strip()
                    right = parts[1].strip()
                    # Only clean if node IDs have invalid characters
                    if re.search(r'[^a-zA-Z0-9_\s]', left) or re.search(r'[^a-zA-Z0-9_\s]', right):
                        left = re.sub(r'[^a-zA-Z0-9_]', '', left)
                        right = re.sub(r'[^a-zA-Z0-9_]', '', right)
                    if left and right:
                        cleaned_lines.append(f"{left} --> {right}")
            elif '[' in line and ']' in line:
                # Node definition - preserve if valid, fix if needed
                match = re.match(r'^([a-zA-Z0-9_]+)\[(.+)\]$', line)
                if match:
                    node_id = match.group(1)
                    label = match.group(2)
                    # Shorten label to be readable (max 18 chars for better readability)
                    label = self._shorten_label(label, max_length=18)
                    # Ensure node_id is short (max 15 chars for readability)
                    if len(node_id) > 15:
                        node_id = node_id[:12] + "_" + str(hash(node_id) % 1000)
                    cleaned_lines.append(f"{node_id}[\"{label}\"]")  # Use quotes for labels with spaces
                else:
                    # Try to fix malformed node definition
                    node_part = line.split('[')[0].strip()
                    node_id = re.sub(r'[^a-zA-Z0-9_]', '', node_part)
                    if node_id:
                        if len(node_id) > 15:
                            node_id = node_id[:12] + "_" + str(hash(node_id) % 1000)
                        label = line.split('[')[1].split(']')[0] if '[' in line and ']' in line else node_id
                        label = self._shorten_label(label, max_length=15)
                        cleaned_lines.append(f"{node_id}[\"{label}\"]")
            elif line and not line.startswith('#'):  # Preserve non-comment lines
                # Try to preserve the line if it looks valid
                cleaned_lines.append(line)
        
        result = "\n".join(cleaned_lines)
        
        # Final validation
        if result and not any(result.strip().lower().startswith(start) for start in valid_starts):
            # If no valid start found, try to add one
            if any('-->' in l for l in cleaned_lines):
                result = "graph TD\n" + result
        
        return result.strip()
    
    def _ensure_readable_labels(self, mermaid_code: str) -> str:
        """Post-process Mermaid code to ensure all labels are readable and short"""
        import re
        lines = mermaid_code.split('\n')
        cleaned_lines = []
        
        for line in lines:
            original_line = line
            line = line.strip()
            if not line:
                continue
            
            # Skip diagram type declarations
            if line.startswith(('graph', 'flowchart', 'sequenceDiagram', 'classDiagram', 'erDiagram', 'gantt', 'pie', 'stateDiagram', 'journey')):
                cleaned_lines.append(line)
                continue
            
            # Skip connection lines (they don't have labels)
            if '-->' in line and '[' not in line:
                cleaned_lines.append(line)
                continue
            
            # Process node definitions with labels: NodeID["Label"] or NodeID[Label] or NodeID(Label)
            # Match quoted labels: NodeID["Label"] or NodeID['Label']
            quoted_pattern = r'^([a-zA-Z0-9_]+)\[(["\'])(.+?)\2\]$'
            quoted_match = re.match(quoted_pattern, line)
            if quoted_match:
                node_id = quoted_match.group(1)
                quote_char = quoted_match.group(2)
                label = quoted_match.group(3)
                
                # Shorten label if too long (stricter limit for readability)
                if len(label) > 18:
                    label = self._shorten_label(label, max_length=18)
                
                # Ensure node_id is short
                if len(node_id) > 15:
                    node_id = node_id[:12] + "_" + str(abs(hash(node_id)) % 1000)
                
                # Use quotes for labels
                cleaned_lines.append(f'{node_id}["{label}"]')
                continue
            
            # Match unquoted labels: NodeID[Label]
            unquoted_pattern = r'^([a-zA-Z0-9_]+)\[(.+?)\]$'
            unquoted_match = re.match(unquoted_pattern, line)
            if unquoted_match:
                node_id = unquoted_match.group(1)
                label = unquoted_match.group(2)
                
                # Remove any existing quotes
                label = label.strip('"\'')
                
                # Shorten label if too long (stricter limit for readability)
                if len(label) > 18:
                    label = self._shorten_label(label, max_length=18)
                
                # Ensure node_id is short
                if len(node_id) > 15:
                    node_id = node_id[:12] + "_" + str(abs(hash(node_id)) % 1000)
                
                # Use quotes for labels
                cleaned_lines.append(f'{node_id}["{label}"]')
                continue
            
            # Keep other lines as-is
            cleaned_lines.append(original_line)
        
        return '\n'.join(cleaned_lines)
    
    def _generate_diagram(self, mermaid_code: str, section_num: str) -> str | None:
        """Generate a diagram image from Mermaid code using multiple services"""
        if not mermaid_code or not mermaid_code.strip():
            logger.warning(f"Empty mermaid code for section {section_num}")
            return None
        
        # Clean the mermaid code
        mermaid_code = self._clean_mermaid_code(mermaid_code)
        if not mermaid_code:
            logger.warning(f"Invalid mermaid code after cleaning for section {section_num}")
            return None
        
        # Post-process to ensure all labels are readable (shorten any remaining long labels)
        mermaid_code = self._ensure_readable_labels(mermaid_code)
        
        # Check for duplicate/similar diagrams - only skip if structure is EXACTLY the same
        normalized_code = self._normalize_mermaid_code(mermaid_code)
        if normalized_code and normalized_code in self.generated_diagrams:
            existing_section = self.generated_diagrams[normalized_code]
            # Only skip if it's a very simple diagram (likely duplicate)
            # For more complex diagrams, allow slight variations
            node_count = len([line for line in mermaid_code.split('\n') if '-->' in line or '[' in line])
            if node_count <= 4:  # Simple diagrams - be strict about duplicates
                logger.debug(f"Skipping duplicate simple diagram for section {section_num} (similar to section {existing_section})")
                return "DUPLICATE"
            else:
                # For complex diagrams, allow if at least 30% different
                logger.debug(f"Complex diagram detected for section {section_num}, allowing even if similar structure")
        
        # Store this diagram to prevent future duplicates
        self.generated_diagrams[normalized_code] = section_num
        
        logger.debug(f"Cleaned mermaid code for section {section_num}: {mermaid_code[:100]}...")
        
        # Try multiple services in order - with improved fallbacks
        # Use a mix: some diagrams via Mermaid, some via PlantUML (for variety and reliability)
        services = [
            ("mermaid.ink", self._try_mermaid_ink),
            ("mermaid.live", self._try_mermaid_live),  # New: Alternative Mermaid renderer
            ("kroki", self._try_kroki),
            ("plantuml", self._try_plantuml),  # PlantUML service for variety
            ("quickchart", self._try_quickchart),
        ]
        
        # Add Python-based fallback only if matplotlib is available
        try:
            import matplotlib
            services.append(("simple_fallback", self._try_simple_diagram_fallback))
        except ImportError:
            logger.debug("matplotlib not available, skipping Python-based diagram fallback")
        
        for service_name, service_func in services:
            try:
                result = service_func(mermaid_code, section_num)
                if result:
                    logger.success(f"Generated diagram using {service_name} for section {section_num}")
                    return result
            except Exception as e:
                logger.warning(f"{service_name} failed for section {section_num}: {e}")
                continue
        
        reason = "All diagram generation services failed"
        logger.error(f"All diagram generation services failed for section {section_num}")
        logger.debug(f"Failed mermaid code: {mermaid_code[:200]}")
        # Remove from tracking if generation failed
        if normalized_code in self.generated_diagrams:
            del self.generated_diagrams[normalized_code]
        # Note: Skipped components tracking is handled in _add_section() where this is called
        return None
    
    def _normalize_mermaid_code(self, mermaid_code: str) -> str:
        """Normalize Mermaid code to detect duplicates/similar diagrams"""
        # Remove whitespace and normalize
        normalized = re.sub(r'\s+', ' ', mermaid_code.strip())
        # Extract just the structure (nodes and edges, ignore labels)
        lines = normalized.split('\n')
        structure_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Extract just node IDs and connections, ignore labels
            if '-->' in line:
                # Edge: extract just node IDs
                parts = line.split('-->')
                if len(parts) == 2:
                    left = re.sub(r'\[.*?\]', '', parts[0].strip())
                    right = re.sub(r'\[.*?\]', '', parts[1].strip())
                    left = re.sub(r'[^a-zA-Z0-9_]', '', left)
                    right = re.sub(r'[^a-zA-Z0-9_]', '', right)
                    structure_lines.append(f"{left}->{right}")
            elif '[' in line and ']' in line:
                # Node definition: extract just node ID
                node_match = re.match(r'^([a-zA-Z0-9_]+)', line)
                if node_match:
                    structure_lines.append(f"node:{node_match.group(1)}")
        # Sort to make structure comparison easier
        structure_lines.sort()
        return '|'.join(structure_lines)
    
    def _try_mermaid_ink(self, mermaid_code: str, section_num: str) -> str | None:
        """Try mermaid.ink service - uses base64 URL encoding"""
        try:
            # mermaid.ink expects base64url encoding (no padding)
            import urllib.parse
            encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii").rstrip('=')
            # URL encode the base64 string
            encoded = urllib.parse.quote(encoded, safe='')
            url = f"https://mermaid.ink/img/{encoded}"
            
            resp = requests.get(url, timeout=20, allow_redirects=True, headers={'User-Agent': 'Mozilla/5.0'})
            
            if resp.status_code == 200 and resp.content and len(resp.content) > 100:
                # Check if it's actually an image (PNG starts with specific bytes)
                if resp.content[:4] == b'\x89PNG':
                    image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
                    image_path = self.output_dir / image_name
                    with open(image_path, "wb") as f:
                        f.write(resp.content)
                    return str(image_path)
                else:
                    # Log what we got for debugging
                    logger.debug(f"mermaid.ink returned non-image (first 100 bytes): {resp.content[:100]}")
                    return None
            else:
                logger.debug(f"mermaid.ink status {resp.status_code}, content length: {len(resp.content) if resp.content else 0}")
                return None
        except Exception as e:
            logger.debug(f"mermaid.ink error: {e}")
            return None
    
    def _try_mermaid_live(self, mermaid_code: str, section_num: str) -> str | None:
        """Try mermaid.live service - alternative Mermaid renderer"""
        try:
            import urllib.parse
            # mermaid.live uses a different encoding approach
            # Try using their API endpoint
            encoded = urllib.parse.quote(mermaid_code, safe='')
            # Alternative: try base64 encoding
            encoded_b64 = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii").rstrip('=')
            encoded_b64 = urllib.parse.quote(encoded_b64, safe='')
            
            # Try multiple possible endpoints
            urls = [
                f"https://mermaid.live/api/png?code={encoded}",
                f"https://mermaid.ink/img/{encoded_b64}",
                f"https://api.mermaid.ink/svg/{encoded_b64}",
            ]
            
            for url in urls:
                try:
                    resp = requests.get(url, timeout=20, allow_redirects=True, headers={'User-Agent': 'Mozilla/5.0'})
                    
                    if resp.status_code == 200 and resp.content and len(resp.content) > 100:
                        # Check if it's actually an image (PNG or SVG)
                        if resp.content[:4] == b'\x89PNG':
                            image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
                            image_path = self.output_dir / image_name
                            with open(image_path, "wb") as f:
                                f.write(resp.content)
                            return str(image_path)
                        elif resp.content[:5] == b'<svg ' or resp.content[:4] == b'<svg':
                            # SVG response - convert to PNG would require additional processing
                            # For now, skip SVG and try next URL
                            continue
                except Exception as e:
                    logger.debug(f"mermaid.live URL {url} failed: {e}")
                    continue
            
            return None
        except Exception as e:
            logger.debug(f"mermaid.live error: {e}")
            return None
    
    def _try_kroki(self, mermaid_code: str, section_num: str) -> str | None:
        """Try Kroki service"""
        try:
            payload = mermaid_code.encode("utf-8")
            compressed = zlib.compress(payload, 9)
            encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
            url = f"https://kroki.io/mermaid/png/{encoded}"
            resp = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
            
            if resp.status_code == 200 and resp.content and len(resp.content) > 100:
                # Verify it's actually a PNG
                if resp.content[:4] == b'\x89PNG':
                    image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
                    image_path = self.output_dir / image_name
                    with open(image_path, "wb") as f:
                        f.write(resp.content)
                    return str(image_path)
            return None
        except Exception as e:
            logger.debug(f"kroki error: {e}")
            return None
    
    def _try_plantuml(self, mermaid_code: str, section_num: str) -> str | None:
        """Try PlantUML service - converts Mermaid to PlantUML and generates diagram"""
        try:
            # Convert simple Mermaid flowcharts to PlantUML
            plantuml_code = self._mermaid_to_plantuml(mermaid_code)
            if not plantuml_code:
                return None
            
            # PlantUML server accepts code via URL encoding
            import urllib.parse
            encoded = urllib.parse.quote(plantuml_code)
            url = f"https://www.plantuml.com/plantuml/png/{encoded}"
            
            resp = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
            
            if resp.status_code == 200 and resp.content and len(resp.content) > 100:
                # Check if it's actually a PNG
                if resp.content[:4] == b'\x89PNG':
                    image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
                    image_path = self.output_dir / image_name
                    with open(image_path, "wb") as f:
                        f.write(resp.content)
                    return str(image_path)
            return None
        except Exception as e:
            logger.debug(f"plantuml error: {e}")
            return None
    
    def _mermaid_to_plantuml(self, mermaid_code: str) -> str | None:
        """Convert simple Mermaid flowchart to PlantUML syntax"""
        try:
            lines = mermaid_code.strip().split('\n')
            if not lines:
                return None
            
            # Check if it's a flowchart/graph
            first_line = lines[0].strip().lower()
            if not (first_line.startswith('graph') or first_line.startswith('flowchart')):
                # PlantUML is best for flowcharts, skip other types
                return None
            
            # Extract direction
            direction = "top to bottom"
            if "LR" in first_line or "left" in first_line:
                direction = "left to right"
            elif "RL" in first_line or "right" in first_line:
                direction = "right to left"
            elif "BT" in first_line or "bottom" in first_line:
                direction = "bottom to top"
            
            plantuml_lines = ["@startuml", ""]
            
            # Extract nodes and edges
            nodes = {}
            edges = []
            
            for line in lines[1:]:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                
                # Extract node: NodeID["Label"] or NodeID[Label]
                node_match = re.match(r'^([a-zA-Z0-9_]+)\[(.+?)\]$', line)
                if node_match:
                    node_id = node_match.group(1)
                    label = node_match.group(2).strip('"\'')
                    # Shorten label if too long
                    if len(label) > 25:
                        label = label[:22] + "..."
                    nodes[node_id] = label
                    continue
                
                # Extract edge: A --> B
                edge_match = re.match(r'^([a-zA-Z0-9_]+)\s*-->\s*([a-zA-Z0-9_]+)', line)
                if edge_match:
                    from_node = edge_match.group(1)
                    to_node = edge_match.group(2)
                    edges.append((from_node, to_node))
                    # Ensure nodes exist
                    if from_node not in nodes:
                        nodes[from_node] = from_node
                    if to_node not in nodes:
                        nodes[to_node] = to_node
            
            if not nodes:
                return None
            
            # Generate PlantUML code
            for node_id, label in nodes.items():
                plantuml_lines.append(f'"{label}" as {node_id}')
            
            plantuml_lines.append("")
            
            for from_node, to_node in edges:
                plantuml_lines.append(f'{from_node} --> {to_node}')
            
            plantuml_lines.append("@enduml")
            
            return '\n'.join(plantuml_lines)
        except Exception as e:
            logger.debug(f"Mermaid to PlantUML conversion error: {e}")
            return None
    
    def _try_quickchart(self, mermaid_code: str, section_num: str) -> str | None:
        """Try QuickChart.io service (supports Mermaid via API)"""
        try:
            # QuickChart doesn't directly support Mermaid, but we can try their chart API
            # For now, skip this and use the simple fallback instead
            return None
        except Exception as e:
            logger.debug(f"quickchart error: {e}")
            return None
    
    def _try_simple_diagram_fallback(self, mermaid_code: str, section_num: str) -> str | None:
        """Generate a simple diagram using Python libraries as final fallback"""
        try:
            import re
            
            # Extract a simple graph structure from mermaid code
            lines = mermaid_code.split('\n')
            nodes = {}
            edges = []
            
            # Find graph type
            graph_type = "TD"  # Top-Down by default
            for line in lines:
                line_lower = line.strip().lower()
                if line_lower.startswith("graph") or line_lower.startswith("flowchart"):
                    if "LR" in line_lower or "left" in line_lower:
                        graph_type = "LR"
                    break
            
            # Extract nodes and edges
            for line in lines:
                line = line.strip()
                if not line or line.startswith(('graph', 'flowchart')):
                    continue
                
                # Extract node definitions: NodeID[Label]
                node_match = re.match(r'^([a-zA-Z0-9_]+)\[(.+)\]$', line)
                if node_match:
                    node_id = node_match.group(1)
                    label = node_match.group(2)[:30]  # Limit label length
                    nodes[node_id] = label
                    continue
                
                # Extract edges: Node1 --> Node2
                edge_match = re.match(r'^([a-zA-Z0-9_]+)\s*-->\s*([a-zA-Z0-9_]+)', line)
                if edge_match:
                    from_node = edge_match.group(1)
                    to_node = edge_match.group(2)
                    edges.append((from_node, to_node))
                    # Ensure nodes exist
                    if from_node not in nodes:
                        nodes[from_node] = from_node[:20]
                    if to_node not in nodes:
                        nodes[to_node] = to_node[:20]
            
            # If we have nodes and edges, create a simple diagram using graphviz or matplotlib
            if nodes and edges:
                return self._create_simple_graphviz_diagram(nodes, edges, section_num, graph_type)
            elif nodes:
                # Just nodes, no edges - create a simple list diagram
                return self._create_simple_node_diagram(nodes, section_num)
            
            return None
        except Exception as e:
            logger.debug(f"Simple diagram fallback error: {e}")
            return None
    
    def _create_simple_graphviz_diagram(self, nodes: dict, edges: list, section_num: str, direction: str = "TD") -> str | None:
        """Create a simple diagram using graphviz"""
        try:
            try:
                import graphviz
            except ImportError:
                logger.debug("graphviz not available (install with: pip install graphviz), trying matplotlib fallback")
                return self._create_simple_matplotlib_diagram(nodes, edges, section_num)
            
            # Create graph
            graph_attr = {'rankdir': 'TB' if direction == "TD" else 'LR', 'size': '8,6', 'dpi': '150'}
            dot = graphviz.Digraph(comment='Diagram', graph_attr=graph_attr)
            dot.attr('node', shape='box', style='rounded,filled', fillcolor='lightblue')
            
            # Add nodes
            for node_id, label in nodes.items():
                dot.node(node_id, label[:25])  # Limit label length
            
            # Add edges
            for from_node, to_node in edges:
                if from_node in nodes and to_node in nodes:
                    dot.edge(from_node, to_node)
            
            # Render to PNG
            image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
            image_path = self.output_dir / image_name
            
            try:
                dot.render(str(image_path).replace('.png', ''), format='png', cleanup=True)
                # graphviz adds .png extension, so adjust path
                actual_path = Path(str(image_path).replace('.png', '') + '.png')
                if actual_path.exists():
                    return str(actual_path)
            except Exception as e:
                logger.debug(f"graphviz render error: {e}")
                return self._create_simple_matplotlib_diagram(nodes, edges, section_num)
            
            return None
        except Exception as e:
            logger.debug(f"graphviz diagram creation error: {e}")
            return self._create_simple_matplotlib_diagram(nodes, edges, section_num)
    
    def _create_simple_matplotlib_diagram(self, nodes: dict, edges: list, section_num: str) -> str | None:
        """Create a simple diagram using matplotlib as final fallback"""
        try:
            try:
                import matplotlib.pyplot as plt
                import matplotlib.patches as mpatches
                from matplotlib.patches import FancyBboxPatch
            except ImportError:
                logger.warning("matplotlib not available (optional dependency). Skipping Python-based diagram generation.")
                return None
            
            fig, ax = plt.subplots(figsize=(10, 8))
            ax.set_xlim(0, 10)
            ax.set_ylim(0, 10)
            ax.axis('off')
            
            # Simple layout: arrange nodes in a grid or tree
            node_positions = {}
            num_nodes = len(nodes)
            
            if num_nodes <= 4:
                # 2x2 grid
                positions = [(2, 8), (8, 8), (2, 4), (8, 4)]
                for i, (node_id, label) in enumerate(nodes.items()):
                    if i < len(positions):
                        node_positions[node_id] = positions[i]
            elif num_nodes <= 6:
                # 3x2 grid
                positions = [(1.5, 8), (5, 8), (8.5, 8), (1.5, 4), (5, 4), (8.5, 4)]
                for i, (node_id, label) in enumerate(nodes.items()):
                    if i < len(positions):
                        node_positions[node_id] = positions[i]
            else:
                # Tree layout
                y_start = 9
                y_step = 1.5
                x_start = 2
                x_step = 1.5
                for i, (node_id, label) in enumerate(nodes.items()):
                    row = i // 3
                    col = i % 3
                    node_positions[node_id] = (x_start + col * x_step, y_start - row * y_step)
            
            # Draw edges first (so they appear behind nodes)
            for from_node, to_node in edges:
                if from_node in node_positions and to_node in node_positions:
                    x1, y1 = node_positions[from_node]
                    x2, y2 = node_positions[to_node]
                    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                              arrowprops=dict(arrowstyle='->', lw=2, color='gray'))
            
            # Draw nodes
            for node_id, (x, y) in node_positions.items():
                label = nodes[node_id][:15]  # Limit label length
                # Draw rounded rectangle
                box = FancyBboxPatch((x-0.8, y-0.3), 1.6, 0.6,
                                    boxstyle="round,pad=0.1",
                                    facecolor='lightblue',
                                    edgecolor='black',
                                    linewidth=1.5)
                ax.add_patch(box)
                ax.text(x, y, label, ha='center', va='center', fontsize=9, weight='bold')
            
            # Save
            image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
            image_path = self.output_dir / image_name
            plt.tight_layout()
            plt.savefig(str(image_path), dpi=150, bbox_inches='tight', facecolor='white')
            plt.close()
            
            return str(image_path)
        except Exception as e:
            logger.debug(f"matplotlib diagram creation error: {e}")
            return None
    
    def _create_simple_node_diagram(self, nodes: dict, section_num: str) -> str | None:
        """Create a simple diagram showing just nodes (no edges)"""
        try:
            try:
                import matplotlib.pyplot as plt
                from matplotlib.patches import FancyBboxPatch
            except ImportError:
                logger.warning("matplotlib not available (optional dependency). Skipping Python-based diagram generation.")
                return None
            
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.set_xlim(0, 10)
            ax.set_ylim(0, 10)
            ax.axis('off')
            
            # Arrange nodes in a simple list
            y_pos = 9
            y_step = 1.2
            x_pos = 5
            
            for i, (node_id, label) in enumerate(nodes.items()):
                if i >= 8:  # Limit to 8 nodes
                    break
                y = y_pos - i * y_step
                label_short = label[:20]
                
                # Draw box
                box = FancyBboxPatch((x_pos-1.5, y-0.3), 3, 0.6,
                                    boxstyle="round,pad=0.1",
                                    facecolor='lightblue',
                                    edgecolor='black',
                                    linewidth=1.5)
                ax.add_patch(box)
                ax.text(x_pos, y, label_short, ha='center', va='center', fontsize=10, weight='bold')
            
            # Save
            image_name = f"diagram_{section_num.replace('.', '_').replace(' ', '_')}.png"
            image_path = self.output_dir / image_name
            plt.tight_layout()
            plt.savefig(str(image_path), dpi=150, bbox_inches='tight', facecolor='white')
            plt.close()
            
            return str(image_path)
        except Exception as e:
            logger.debug(f"Simple node diagram creation error: {e}")
            return None
    
    def _add_references_section(self, doc: DOCXGenerator, content: Dict[str, Any], outline: Dict[str, Any], codebase_info: Dict[str, Any] = None):
        """Add references section with AI-generated references"""
        logger.info("Adding References section...")
        
        # Generate references using AI
        from utils.llm_client import llm_client
        from config.prompts import GENERATE_REFERENCES_PROMPT
        
        project_name = content.get("report_title", "Technical Report")
        technologies = codebase_info.get("technologies", []) if codebase_info else []
        
        # Extract additional technologies from content if not provided
        if not technologies:
            chapters = content.get("chapters", [])
            for chapter in chapters:
                chapter_title = chapter.get("chapter_title", "")
                sections = chapter.get("sections", [])
                for section in sections:
                    section_title = section.get("title", "")
                    content_text = section.get("content", "").lower()
                    # Collect relevant keywords
                    if "python" in section_title.lower() or "python" in chapter_title.lower() or "python" in content_text:
                        technologies.append("Python")
                    if "api" in section_title.lower() or "api" in chapter_title.lower() or "api" in content_text:
                        technologies.append("API")
                    if "database" in section_title.lower() or "database" in content_text:
                        technologies.append("Database")
        
        tech_string = ", ".join(set(technologies)) if technologies else "Software Development"
        
        prompt = GENERATE_REFERENCES_PROMPT.format(
            project_title=project_name,
            technologies=tech_string
        )
        
        try:
            references = llm_client.generate_json(
                prompt=prompt,
                system_prompt="You are a technical writer. Generate academic and technical references in proper citation format.",
                temperature=0.3,
                timeout=300
            )
            
            # Add references heading using DOCXGenerator method (centered, properly formatted)
            doc.add_conclusion_heading("REFERENCES")
            
            # Add references
            ref_list = references.get("references", [])
            
            # CRITICAL FIX: Handle malformed JSON where ref_list might be a string or nested structure
            if isinstance(ref_list, str):
                # If it's a string, try to split by newlines or treat as single reference
                logger.warning("References is a string, attempting to parse...")
                ref_list = [ref_list] if ref_list.strip() else []
            elif not isinstance(ref_list, list):
                # If it's not a list, try to convert
                logger.warning(f"References is not a list (type: {type(ref_list)}), attempting to convert...")
                ref_list = []
            
            if ref_list:
                # Ensure ref_list is a list of strings (handle dicts, nested structures, etc.)
                formatted_refs = []
                for ref in ref_list:
                    # Handle deeply nested structures (common in malformed JSON)
                    if isinstance(ref, dict):
                        # Try to extract text from various possible keys
                        ref_text = (ref.get("text") or ref.get("reference") or 
                                   ref.get("citation") or ref.get("title") or 
                                   str(ref))
                        # If still a dict, try to get first value
                        if isinstance(ref_text, dict):
                            ref_text = list(ref_text.values())[0] if ref_text else str(ref)
                    elif isinstance(ref, str):
                        ref_text = ref.strip()
                    elif isinstance(ref, list):
                        # If it's a list, take first element
                        ref_text = str(ref[0]) if ref else ""
                    else:
                        ref_text = str(ref).strip()
                    
                    # Only add non-empty references
                    if ref_text and len(ref_text) > 10:  # Minimum length check
                        formatted_refs.append(ref_text)
                
                # HARDENING: Validate reference scope against content metadata
                # Create metadata dict for validation
                content_metadata = {
                    "report_title": content.get("report_title", ""),
                    "chapters": content.get("chapters", [])
                }
                formatted_refs = self._validate_reference_scope(formatted_refs, content_metadata)
                
                # Add each reference with proper formatting
                for i, ref_text in enumerate(formatted_refs, 1):
                    p = doc.document.add_paragraph()
                    # Use hanging indent format for references
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.hanging_indent = Inches(-0.5)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add reference number and text
                    run = p.add_run(f"[{i}] ")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    
                    ref_run = p.add_run(ref_text)
                    ref_run.font.size = Pt(11)
                    ref_run.font.bold = False
            else:
                # Fallback if AI fails
                doc.add_paragraph("References will be added here.")
            
            logger.success(f"References section added with {len(ref_list)} references")
        except Exception as e:
            logger.warning(f"Failed to generate references: {e}. Adding placeholder.")
            doc.add_conclusion_heading("REFERENCES")
            doc.add_paragraph("References will be added here.")
    
    def _add_list_of_figures(self, doc: DOCXGenerator):
        """Add List of Figures section"""
        if not hasattr(doc, 'figures_list') or not doc.figures_list:
            return
        
        logger.info("Adding List of Figures...")
        
        # Add heading (centered, bold, uppercase)
        heading = doc.document.add_heading("LIST OF FIGURES", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if heading.runs:
            heading.runs[0].font.bold = True
        else:
            run = heading.add_run("LIST OF FIGURES")
            run.font.bold = True
        
        # Add each figure entry
        for fig in doc.figures_list:
            p = doc.document.add_paragraph()
            # Format: "Fig. 3.1. System Architecture Flowchart"
            label = fig.get("label", "").replace("Fig ", "Fig. ").replace("Fig.", "Fig.")
            if not label.startswith("Fig."):
                label = f"Fig. {label}" if label else "Fig."
            description = fig.get("description", "")
            
            run = p.add_run(f"{label} {description}")
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
        
        logger.success(f"List of Figures added with {len(doc.figures_list)} entries")
    
    def _add_list_of_tables(self, doc: DOCXGenerator):
        """Add List of Tables section"""
        if not hasattr(doc, 'tables_list') or not doc.tables_list:
            return
        
        logger.info("Adding List of Tables...")
        
        # Add heading (centered, bold, uppercase)
        heading = doc.document.add_heading("LIST OF TABLES", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if heading.runs:
            heading.runs[0].font.bold = True
        else:
            run = heading.add_run("LIST OF TABLES")
            run.font.bold = True
        
        # Add each table entry
        for table in doc.tables_list:
            p = doc.document.add_paragraph()
            # Format: "Table 3.1. Semester Course table"
            label = table.get("label", "").replace("Table ", "Table ")
            if not label.startswith("Table"):
                label = f"Table {label}" if label else "Table"
            title = table.get("title", "")
            
            run = p.add_run(f"{label} {title}")
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
        
        logger.success(f"List of Tables added with {len(doc.tables_list)} entries")
    
    # ============================================================
    # HARDENING METHODS - Final Quality Gate (Domain-Agnostic)
    # ============================================================
    
    def _sanitize_text_coherence(self, text: str) -> str:
        """
        Final text coherence sanitation pass.
        Removes unmatched quotes, duplicated phrases, dangling conjunctions,
        and fixes objectively broken grammar.
        
        This is CLEANUP ONLY, not rewriting.
        """
        if not text or not isinstance(text, str):
            return text
        
        # Remove unmatched quotes (fix common issues)
        # Count quotes and balance them
        single_quotes = text.count("'") - text.count("\\'")
        double_quotes = text.count('"') - text.count('\\"')
        
        # If odd number of quotes, try to fix common patterns
        if single_quotes % 2 == 1:
            # Remove standalone unmatched single quotes (fixed-width lookbehind only)
            text = re.sub(r"(?<!\w)'(?=\s|$)", "", text)  # Quote before space/end
            text = re.sub(r"^\s*'", "", text)  # Quote at start of line/string
            text = re.sub(r"\s+'(?!\w)", " ", text)  # Quote after space (not followed by word)
        
        if double_quotes % 2 == 1:
            # Remove standalone unmatched double quotes (fixed-width lookbehind only)
            text = re.sub(r'(?<!\w)"(?=\s|$)', '', text)  # Quote before space/end
            text = re.sub(r'^\s*"', '', text)  # Quote at start of line/string
            text = re.sub(r'\s+"(?!\w)', ' ', text)  # Quote after space (not followed by word)
        
        # Remove duplicated phrases (consecutive identical sentences)
        sentences = re.split(r'([.!?]\s+)', text)
        cleaned_sentences = []
        prev_sentence = ""
        for i in range(0, len(sentences) - 1, 2):
            if i + 1 < len(sentences):
                sentence = sentences[i] + sentences[i + 1]
                sentence_stripped = sentence.strip()
                # Skip if identical to previous (allow slight variations)
                if sentence_stripped.lower() != prev_sentence.lower() or len(cleaned_sentences) == 0:
                    cleaned_sentences.append(sentence)
                    prev_sentence = sentence_stripped
            else:
                cleaned_sentences.append(sentences[i])
        text = ''.join(cleaned_sentences) if cleaned_sentences else text
        
        # Remove dangling conjunctions at sentence start (after cleanup)
        text = re.sub(r'^\s*(And|But|Or|However|Moreover|Furthermore|Additionally|Also|Yet)\s+[.,]', '', text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Fix broken grammar: incomplete sentences (sentences without ending punctuation)
        # Only fix if it's clearly broken (ends with comma or no punctuation before new paragraph)
        paragraphs = text.split('\n\n')
        fixed_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if para:
                # If paragraph ends with comma and no sentence-ending punctuation, add period
                if para[-1] == ',' and '.' not in para[-50:]:
                    para = para.rstrip(',') + '.'
                # If paragraph has no ending punctuation at all, add period
                elif para and para[-1] not in '.!?':
                    if len(para) > 20:  # Only for substantial paragraphs
                        para = para + '.'
                fixed_paragraphs.append(para)
        text = '\n\n'.join(fixed_paragraphs)
        
        # Remove excessive whitespace
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs
        text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines
        text = text.strip()
        
        return text
    
    def _enforce_epistemic_boundary(self, text: str) -> str:
        """
        Enforce epistemic boundary: Remove/neutralize consequence/outcome/runtime
        claims that lack explicit evidence in input data.
        
        Universal rule: If statement describes consequences/outcomes/failures/runtime
        behavior AND no explicit evidence provided, replace with absence-based description.
        """
        if not text or not isinstance(text, str):
            return text
        
        # Pattern: Claims about consequences/outcomes that are not provable from static code analysis
        # These patterns are domain-agnostic and catch universal consequence language
        
        consequence_patterns = [
            # "will cause", "would cause", "may cause" - without evidence
            (r'\b(?:will|would|may|might|can|could)\s+cause\s+(?:runtime\s+)?(?:errors?|exceptions?|failures?|crashes?|issues?)', 
             'does not include logic to prevent'),
            # "results in", "leads to" - outcome claims
            (r'\b(?:results?\s+in|leads?\s+to)\s+(?:runtime\s+)?(?:errors?|exceptions?|failures?|crashes?)', 
             'does not include safeguards to prevent'),
            # "throws", "throws an exception" - runtime behavior
            (r'\bthrows?\s+(?:an?\s+)?(?:unhandled\s+)?(?:runtime\s+)?exceptions?', 
             'does not include exception handling'),
            # "crashes", "will crash" - runtime outcome
            (r'\b(?:will|would|may|might)?\s*crashes?\s+(?:on|when|if)', 
             'does not include logic to prevent crashes'),
            # "produces incorrect results" - outcome claim
            (r'\bproduces?\s+(?:incorrect|invalid|wrong|NaN|Infinity)', 
             'does not include validation to ensure correct'),
            # "fails silently" - runtime behavior
            (r'\bfails?\s+silently', 
             'does not provide error feedback'),
            # "will fail", "may fail" - outcome prediction
            (r'\b(?:will|would|may|might)\s+fail\s+(?:to|when|if)', 
             'does not include logic to handle'),
        ]
        
        for pattern, replacement in consequence_patterns:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Remove standalone consequence verbs in error contexts
        if re.search(r'(?:error|exception|failure|crash|invalid)', text, re.IGNORECASE):
            standalone_consequences = [
                (r'\bcause\s+(?:runtime\s+)?errors?', 'does not include logic to prevent errors'),
                (r'\bproduces?\s+NaN', 'does not include validation to prevent NaN'),
                (r'\bproduces?\s+Infinity', 'does not include validation to prevent Infinity'),
            ]
            for pattern, replacement in standalone_consequences:
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _validate_heading(self, heading_text: str) -> str:
        """
        Validate and sanitize heading text.
        Strip malformed characters, validate format.
        Returns sanitized heading or empty string if too corrupted.
        """
        if not heading_text or not isinstance(heading_text, str):
            return ""
        
        # Strip leading/trailing whitespace
        heading_text = heading_text.strip()
        
        # Remove control characters (except newlines for multi-line headings)
        heading_text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', heading_text)
        
        # Remove excessive special characters (more than 2 consecutive)
        heading_text = re.sub(r'[^\w\s-]{3,}', '', heading_text)
        
        # Remove excessive whitespace
        heading_text = re.sub(r'\s+', ' ', heading_text)
        heading_text = heading_text.strip()
        
        # If heading is too short or only special chars, return empty (will use fallback)
        if len(heading_text) < 2 or not re.search(r'[a-zA-Z0-9]', heading_text):
            logger.warning(f"Heading too corrupted, using fallback")
            return ""
        
        return heading_text
    
    def _validate_table_cell(self, cell_text: str) -> str:
        """
        Validate and sanitize a single table cell.
        Applies text coherence and epistemic boundary enforcement.
        """
        if not cell_text or not isinstance(cell_text, str):
            return str(cell_text) if cell_text else ""
        
        # Apply coherence sanitation
        sanitized = self._sanitize_text_coherence(cell_text)
        
        # Apply epistemic boundary enforcement
        sanitized = self._enforce_epistemic_boundary(sanitized)
        
        return sanitized
    
    def _validate_table_data(self, table_data: list) -> list:
        """
        Validate and sanitize entire table.
        Processes every cell through validation.
        """
        if not table_data or not isinstance(table_data, list):
            return []
        
        validated_table = []
        for row in table_data:
            if not isinstance(row, list):
                continue
            validated_row = [self._validate_table_cell(str(cell)) for cell in row]
            validated_table.append(validated_row)
        
        return validated_table if validated_table else []
    
    def _validate_diagram_semantics(self, mermaid_code: str) -> tuple[str, bool]:
        """
        Validate diagram semantics: Ensure diagram encodes structure/flow,
        not outcomes/errors/runtime behavior.
        
        Returns: (validated_code, is_valid)
        If invalid semantics found, returns neutralized version or empty string.
        """
        if not mermaid_code or not isinstance(mermaid_code, str):
            return ("", False)
        
        lines = mermaid_code.strip().split('\n')
        if not lines:
            return ("", False)
        
        # Patterns that indicate outcome/error/runtime semantics (domain-agnostic)
        invalid_patterns = [
            r'(?:error|exception|failure|crash|invalid|NaN|Infinity|fail|throw)',
            r'(?:will\s+cause|would\s+cause|may\s+cause|results?\s+in|leads?\s+to)',
            r'(?:runtime\s+error|runtime\s+exception)',
        ]
        
        # Check if diagram contains invalid semantic patterns
        code_lower = mermaid_code.lower()
        has_invalid_semantics = any(re.search(pattern, code_lower, re.IGNORECASE) for pattern in invalid_patterns)
        
        if has_invalid_semantics:
            # Neutralize: Replace outcome nodes with neutral terminal
            # This is a conservative approach - we modify the diagram structure
            neutralized_lines = []
            for line in lines:
                line_lower = line.lower()
                # If line contains invalid semantics, try to neutralize it
                if any(re.search(pattern, line_lower) for pattern in invalid_patterns):
                    # Replace with neutral node: "No explicit handling logic present"
                    # Only replace node definitions, not connections
                    if '[' in line and ']' in line:
                        # Extract node ID
                        node_match = re.match(r'^([a-zA-Z0-9_]+)\[', line)
                        if node_match:
                            node_id = node_match.group(1)
                            neutralized_lines.append(f'{node_id}["No explicit handling logic present"]')
                        else:
                            # Skip this line (remove invalid node)
                            continue
                    elif '-->' in line:
                        # Skip connections involving invalid nodes
                        continue
                    else:
                        neutralized_lines.append(line)
                else:
                    neutralized_lines.append(line)
            
            neutralized_code = '\n'.join(neutralized_lines)
            if neutralized_code.strip():
                logger.warning("Diagram contained outcome/error semantics, neutralized")
                return (neutralized_code, True)  # Return neutralized version
            else:
                logger.warning("Diagram was entirely outcome-based, rejected")
                return ("", False)  # Reject entirely
        
        return (mermaid_code, True)  # Valid structure/flow diagram
    
    def _validate_reference_scope(self, references: list, content_metadata: Dict[str, Any]) -> list:
        """
        Validate references against input metadata.
        Drop overly generic, unrelated, or speculative references.
        
        References must align with detected technologies/concepts from content.
        """
        if not references or not isinstance(references, list):
            return []
        
        # Extract technologies/concepts from metadata (generic approach)
        technologies = set()
        if content_metadata:
            # Check for common technology keywords (domain-agnostic)
            content_text = json.dumps(content_metadata).lower()
            tech_keywords = ['python', 'javascript', 'java', 'c++', 'c#', 'sql', 'database', 
                           'api', 'rest', 'http', 'web', 'frontend', 'backend', 'server',
                           'framework', 'library', 'tool', 'system', 'application']
            
            for keyword in tech_keywords:
                if keyword in content_text:
                    technologies.add(keyword)
        
        validated_refs = []
        for ref in references:
            if not isinstance(ref, str):
                continue
            
            ref_lower = ref.lower()
            
            # Skip if reference is too generic/unrelated
            # Generic patterns that indicate non-specific references
            generic_patterns = [
                r'^(general|introduction|overview|basics)\s+(?:to|of|about)',
                r'^(how\s+to|tutorial|guide)\s+(?:for|on)',
            ]
            
            is_too_generic = any(re.search(pattern, ref_lower) for pattern in generic_patterns)
            
            # If we have technologies detected, check if reference aligns
            if technologies:
                has_relevance = any(tech in ref_lower for tech in technologies)
                if not has_relevance and is_too_generic:
                    logger.debug(f"Dropping generic/unrelated reference: {ref[:50]}...")
                    continue
            
            # Apply text coherence to reference
            sanitized_ref = self._sanitize_text_coherence(ref)
            if sanitized_ref and len(sanitized_ref) > 10:
                validated_refs.append(sanitized_ref)
        
        # Scale reference count with document size (generic heuristic)
        # More chapters/sections = more references allowed
        if content_metadata and 'chapters' in str(content_metadata):
            # Rough heuristic: 2-3 references per chapter
            max_refs = len(re.findall(r'"chapter', json.dumps(content_metadata), re.IGNORECASE)) * 3
            if max_refs > 0:
                validated_refs = validated_refs[:max_refs]
        
        return validated_refs
    
    def _log_skipped_components_summary(self):
        """Log summary of skipped/failed components for user awareness"""
        # Deduplicate entries (same section + reason combination)
        unique_diagrams = list(set(self.skipped_components["diagrams"]))
        unique_tables = list(set(self.skipped_components["tables"]))
        unique_subsections = list(set(self.skipped_components["subsections"]))
        
        total_skipped = len(unique_diagrams) + len(unique_tables) + len(unique_subsections)
        
        if total_skipped > 0:
            logger.info("=" * 60)
            logger.info("SKIPPED COMPONENTS SUMMARY")
            logger.info("=" * 60)
            
            if unique_diagrams:
                logger.warning(f"Skipped {len(unique_diagrams)} diagram(s):")
                for section_num, reason in unique_diagrams:
                    logger.warning(f"  - Section {section_num}: {reason}")
            
            if unique_tables:
                logger.warning(f"Skipped {len(unique_tables)} table(s):")
                for section_num, reason in unique_tables:
                    logger.warning(f"  - Section {section_num}: {reason}")
            
            if unique_subsections:
                logger.warning(f"Skipped {len(unique_subsections)} subsection(s):")
                for section_num, reason in unique_subsections:
                    logger.warning(f"  - Section {section_num}: {reason}")
            
            logger.info("=" * 60)

    def run(
        self,
        content: Dict[str, Any],
        guidelines: Dict[str, Any],
        outline: Dict[str, Any]
    ) -> str:
        """
        Run complete building pipeline
        
        Returns:
            Path to generated document
        """
        return self.build_document(content, guidelines, outline)


# Convenience function
def build_docx_report(
    content: Dict[str, Any],
    guidelines: Dict[str, Any],
    outline: Dict[str, Any],
    output_filename: str = "Technical_Report.docx",
    job_id: str = None
) -> str:
    """
    Build DOCX report from components
    
    Args:
        content: Chapter content
        guidelines: Formatting guidelines
        outline: Report structure
        output_filename: Output file name
        job_id: Optional job ID for isolated output directory
        
    Returns:
        Path to generated document
    """
    agent = BuilderAgent(job_id=job_id)
    return agent.build_document(content, guidelines, outline, output_filename)