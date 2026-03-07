"""
DOCX Generator - Creates formatted Word documents
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Any, Optional, List
from loguru import logger


class DOCXGenerator:
    """Generate formatted DOCX documents based on guidelines"""
    
    def __init__(self, guidelines: Dict[str, Any]):
        self.guidelines = guidelines
        self.document = Document()
        self._setup_page()
        self._setup_styles()
    
    def _setup_page(self):
        """Configure page setup (margins, size)"""
        sections = self.document.sections
        for section in sections:
            # Set margins
            margins = self.guidelines.get("page_setup", {}).get("margins", {})
            section.top_margin = self._parse_measurement(margins.get("top", "1.0in"))
            section.bottom_margin = self._parse_measurement(margins.get("bottom", "1.0in"))
            section.left_margin = self._parse_measurement(margins.get("left", "1.25in"))
            section.right_margin = self._parse_measurement(margins.get("right", "1.0in"))
            
            # Set paper size (A4)
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
        
        logger.info("Page setup configured")
    
    def _parse_measurement(self, value: str) -> int:
        """Convert measurement string to docx units (twips)"""
        # Handle dict input (shouldn't happen, but be safe)
        if isinstance(value, dict):
            logger.warning(f"Received dict instead of string for measurement: {value}, defaulting to 1.0 inches")
            return Inches(1.0)
        
        if isinstance(value, (int, float)):
            return Inches(float(value))
        
        if value is None:
            logger.warning("Received None for measurement, defaulting to 1.0 inches")
            return Inches(1.0)
        
        value = str(value).lower().strip()
        
        # Extract numeric value using regex to handle various formats
        # Try to find a number (integer or float) in the string
        number_match = re.search(r'(\d+\.?\d*)', value)
        if not number_match:
            # If no number found, default to 1.0 inches
            logger.warning(f"Could not parse measurement '{value}', defaulting to 1.0 inches")
            return Inches(1.0)
        
        numeric_value = float(number_match.group(1))
        
        # Check for unit indicators
        # Check for "inches" (full word) or "in" as a standalone unit (not part of another word)
        has_inches = "inches" in value or re.search(r'\bin\b', value) or value.endswith("in")
        has_cm = "centimeters" in value or "cm" in value
        
        if has_inches:
            return Inches(numeric_value)
        elif has_cm:
            return Inches(numeric_value / 2.54)
        else:
            # Default to inches if no unit specified
            return Inches(numeric_value)
    
    def _setup_styles(self):
        """Configure text styles based on guidelines"""
        styles = self.document.styles
        fonts = self.guidelines.get("fonts", {})
        spacing = self.guidelines.get("spacing", {})
        typography = self.guidelines.get("typography", {})
        
        # Normal style (body text)
        normal = styles['Normal']
        font = normal.font
        
        # Get font family - check both fonts and typography
        font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
        font.name = font_family
        
        # Get body text size - handle both dict and direct value
        body_text_config = fonts.get("body_text") or typography.get("body_text")
        if isinstance(body_text_config, dict):
            font.size = Pt(body_text_config.get("size", 12))
        else:
            font.size = Pt(12)
        
        # Paragraph formatting
        paragraph_format = normal.paragraph_format
        
        # Handle line_spacing - can be a number or a dict
        line_spacing_value = spacing.get("line_spacing", 1.5)
        if isinstance(line_spacing_value, dict):
            # If it's a dict, use regular_text value
            line_spacing_value = line_spacing_value.get("regular_text", 1.5)
        paragraph_format.line_spacing = line_spacing_value
        
        # Handle paragraph spacing - check if they exist as direct values or in nested structure
        space_before = spacing.get("paragraph_spacing_before", 0)
        if isinstance(space_before, dict):
            space_before = 0
        paragraph_format.space_before = Pt(space_before)
        
        space_after = spacing.get("paragraph_spacing_after", 12)
        if isinstance(space_after, dict):
            space_after = 12
        paragraph_format.space_after = Pt(space_after)
        
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        logger.info("Document styles configured")
    
    def add_chapter_heading(self, chapter_num: int, title: str):
        """Add chapter heading (centered)"""
        fonts = self.guidelines.get("fonts", {})
        typography = self.guidelines.get("typography", {})
        chapter_font = fonts.get("chapter_heading") or typography.get("chapter_heading", {})
        
        heading = self.document.add_heading(level=1)
        run = heading.add_run(f"CHAPTER {chapter_num}. {title.upper()}")
        
        font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
        run.font.name = font_family
        run.font.size = Pt(chapter_font.get("size", 16) if isinstance(chapter_font, dict) else 16)
        run.font.bold = chapter_font.get("bold", True) if isinstance(chapter_font, dict) else True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center chapter headings
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
        
        logger.debug(f"Added chapter: {chapter_num}. {title}")

    def add_conclusion_heading(self, title: str):
        """Add conclusion heading without chapter number label (centered)"""
        fonts = self.guidelines.get("fonts", {})
        typography = self.guidelines.get("typography", {})
        chapter_font = fonts.get("chapter_heading") or typography.get("chapter_heading", {})

        heading = self.document.add_heading(level=1)
        run = heading.add_run(title.upper())

        font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
        run.font.name = font_family
        run.font.size = Pt(chapter_font.get("size", 16) if isinstance(chapter_font, dict) else 16)
        run.font.bold = chapter_font.get("bold", True) if isinstance(chapter_font, dict) else True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center conclusion heading
        run.font.color.rgb = RGBColor(0, 0, 0)

        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)

        logger.debug(f"Added conclusion heading: {title}")
    
    def add_section_heading(self, section_num: str, title: str, is_conclusion: bool = False):
        """Add section heading"""
        fonts = self.guidelines.get("fonts", {})
        typography = self.guidelines.get("typography", {})
        section_font = fonts.get("section_heading") or typography.get("section_heading", {})
        
        heading = self.document.add_heading(level=2)
        if is_conclusion or not section_num:
            text = title.upper()
        else:
            text = f"{section_num} {title.upper()}"
        run = heading.add_run(text)
        
        font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
        run.font.name = font_family
        run.font.size = Pt(section_font.get("size", 14) if isinstance(section_font, dict) else 14)
        run.font.bold = section_font.get("bold", True) if isinstance(section_font, dict) else True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
        
        logger.debug(f"Added section: {section_num} {title}")
    
    def add_subsection_heading(self, subsection_num: str, title: str):
        """Add subsection heading"""
        fonts = self.guidelines.get("fonts", {})
        typography = self.guidelines.get("typography", {})
        subsection_font = fonts.get("subsection_heading") or typography.get("subsection_heading", {})
        
        heading = self.document.add_heading(level=3)
        # Only capitalize first letter for subsections
        run = heading.add_run(f"{subsection_num} {title}")
        
        font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
        run.font.name = font_family
        run.font.size = Pt(subsection_font.get("size", 12) if isinstance(subsection_font, dict) else 12)
        run.font.bold = subsection_font.get("bold", True) if isinstance(subsection_font, dict) else True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    def add_paragraph(self, text: str):
        """Add body paragraph"""
        if not text.strip():
            return
        
        paragraph = self.document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Apply font to all runs
        fonts = self.guidelines.get("fonts", {})
        typography = self.guidelines.get("typography", {})
        for run in paragraph.runs:
            font_family = fonts.get("family") or typography.get("font_family") or "Times New Roman"
            run.font.name = font_family
            body_text_config = fonts.get("body_text") or typography.get("body_text")
            if isinstance(body_text_config, dict):
                run.font.size = Pt(body_text_config.get("size", 12))
            else:
                run.font.size = Pt(12)
    
    def add_page_break(self):
        """Add page break"""
        self.document.add_page_break()
    
    def add_table_of_contents(self, outline: Dict[str, Any]):
        """Add Word-native TOC field (user can update via References > Table of Contents)"""
        heading = self.document.add_heading("Table of Contents", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.document.add_paragraph()
        run = p.add_run()

        # Begin field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        # Instruction text for TOC (levels 1-3, hyperlinks, hide tab/page numbers in web view)
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr_text)

        # Separate field
        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        run._r.append(fld_char_separate)

        # Placeholder text that Word replaces on update
        placeholder = run.add_text("Table of Contents will be generated. Right-click and select 'Update Field' to refresh.")

        # End field
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)

        # Don't add page break after TOC - let it flow naturally
        logger.info("Inserted Word TOC field (update via References > Table of Contents).")
    
    def _add_static_toc(self, outline: Dict[str, Any]):
        """Generate a static table of contents from the outline"""
        chapters = outline.get("chapters", [])
        
        for chapter in chapters:
            chapter_num = chapter.get("number", 1)
            chapter_title = chapter.get("title", "Chapter")
            is_conclusion = "conclusion" in chapter_title.lower()
            
            # Chapter entry
            p = self.document.add_paragraph()
            if is_conclusion:
                run = p.add_run(chapter_title.upper())
            else:
                run = p.add_run(f"CHAPTER {chapter_num}. {chapter_title.upper()}")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)
            
            # Sections
            sections = chapter.get("sections", [])
            for section in sections:
                section_num = section.get("number", "")
                section_title = section.get("title", "Section")
                
                # Only show section numbers for non-conclusion chapters
                if is_conclusion or not section_num:
                    toc_text = section_title
                else:
                    toc_text = f"{section_num} {section_title}"
                
                p = self.document.add_paragraph()
                run = p.add_run(toc_text)
                run.font.size = Pt(11)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)
                
                # Subsections
                subsections = section.get("subsections", [])
                for subsection in subsections:
                    sub_num = subsection.get("number", "")
                    sub_title = subsection.get("title", "Subsection")
                    
                    p = self.document.add_paragraph()
                    run = p.add_run(f"{sub_num} {sub_title}")
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(1.0)
                    p.paragraph_format.space_after = Pt(2)
        
        logger.debug("Static TOC generated")

    def add_table(self, data: List[List[str]], title: Optional[str] = None, table_label: Optional[str] = None):
        """Add a formatted table with label"""
        if not data or not data[0]:
            return
        
        # Initialize tables_list if not exists
        if not hasattr(self, 'tables_list'):
            self.tables_list = []
        
        # Add title/label if provided
        if table_label:
            # Format label (e.g., "Table 3.1")
            formatted_label = table_label.replace("Table ", "Table ").replace("Table.", "Table")
            if not formatted_label.startswith("Table"):
                formatted_label = f"Table {formatted_label}" if formatted_label else "Table"
            
            title_p = self.document.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            display_text = f"{formatted_label}"
            if title:
                display_text += f" - {title}"
            run = title_p.add_run(display_text)
            run.font.bold = True
            run.font.size = Pt(11)
            
            # Track table for List of Tables
            self.tables_list.append({
                "label": formatted_label,
                "title": title or "Table"
            })
        elif title:
            title_p = self.document.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run(f"Table: {title}")
            run.font.bold = True
            run.font.size = Pt(11)

        table = self.document.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        
        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = str(cell_text)
                # Formatting first row as header
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            
        self.document.add_paragraph("") # Space after table

    def add_figure(self, image_path: str, label: str, description: str):
        """Add a real image/figure with label"""
        try:
            # Initialize figures_list if not exists
            if not hasattr(self, 'figures_list'):
                self.figures_list = []
            
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # Use a moderate width so the diagram does not dominate the entire page
            run.add_picture(image_path, width=Inches(4.0))
            
            # Format label (e.g., "Fig. 3.1" or "Fig 3.1")
            formatted_label = label.replace("Fig ", "Fig. ").replace("Fig.", "Fig.")
            if not formatted_label.startswith("Fig."):
                formatted_label = f"Fig. {formatted_label}" if formatted_label else "Fig."
            
            # Label
            label_p = self.document.add_paragraph()
            label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = label_p.add_run(f"{formatted_label} - {description}")
            run.font.italic = True
            run.font.size = Pt(10)
            
            # Track figure for List of Figures
            self.figures_list.append({
                "label": formatted_label,
                "description": description
            })
            
            self.document.add_paragraph("") # Space after
            logger.info(f"Added figure: {label}")
        except Exception as e:
            logger.error(f"Failed to add image {image_path}: {e}")
            self.add_figure_placeholder(label, description)

    def add_figure_placeholder(self, label: str, description: str):
        """Add a placeholder for a figure/diagram"""
        # Create a frame-like paragraph
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Box placeholder
        run = p.add_run("\n[ DIAGRAM / FIGURE PLACEHOLDER ]\n")
        run.bold = True
        
        # Label
        label_p = self.document.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label_p.add_run(f"Figure: {label} - {description}")
        run.font.italic = True
        run.font.size = Pt(10)
        
        self.document.add_paragraph("") # Space after
    
    def add_header_footer(self, student_id: str = "22AIML056"):
        """Add headers and footers with page numbers"""
        for section in self.document.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = ""
            
            # Set up tab stops for header (Left, Center, Right)
            header_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            
            run_left = header_para.add_run(student_id)
            run_left.font.size = Pt(10)
            
            header_para.add_run("\t") # Center tab
            
            header_para.add_run("\t") # Right tab
            run_right = header_para.add_run("Technical Documentation")
            run_right.font.size = Pt(10)
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = ""
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Footer content
            hf_config = self.guidelines.get("headers_footers", {})
            footer_left = hf_config.get("footer_left", "CSPIT")
            
            run = footer_para.add_run(f"{footer_left} | Page ")
            run.font.size = Pt(10)
            
            # Add page number
            self._add_page_number(footer_para)
            
            logger.info("Headers and footers configured for section")
    
    def _add_page_number(self, paragraph):
        """Add page number field to paragraph"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def save(self, output_path: str):
        """Save document to file"""
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        self.document.save(str(output_file))
        logger.success(f"Document saved: {output_file}")
        return str(output_file)


def create_document(guidelines: Dict[str, Any]) -> DOCXGenerator:
    """
    Create a new DOCX document with formatting
    
    Args:
        guidelines: Formatting guidelines
        
    Returns:
        DOCXGenerator instance
    """
    return DOCXGenerator(guidelines)